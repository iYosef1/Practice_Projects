from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Inches, Mm, Cm, Pt, Twips, Emu
from docx.api import Document
from math import ceil

#document = Document()


def page_orientation_size(document, size, pg_no, landscape = False):
    """
    This function readjusts the size and orientation of the document and its page_no
    passed in. The size parameter requires a tuple (unit, height, width). The unit
    parameter must be one of the following arguments: Inches, Mm, Cm, Pt, Twips, Emu.
    Ensure you have passed in an argument for both the landscape and size parameters
    even if only one adjustment is requested to be made to the page.
    """
    if pg_no in list(range(len(document.sections))):
        pass
    else:
        print('The arguement passed into pg_no is not valid')
        return
    current_section = document.sections[pg_no]
    if isinstance(size, tuple):
        current_section.page_height = size[0](size[1])
        current_section.page_width = size[0](size[2])
    else:
        print('The size parameter requires a tuple argument (unit, height, width)')
    if landscape == True:
        current_section.orientation = WD_ORIENT.LANDSCAPE
        landscape_height = current_section.page_width
        landscape_width = current_section.page_height
        current_section.page_height = Inches(landscape_height.inches)
        current_section.page_width = Inches(landscape_width.inches)

#page_orientation_size(document, (Inches, 11.7, 16.5), 0, False)






def change_margins(document, dimensions, pg_no):
    """
    This function changes the margins of the page number passed into the
    parameter, pg_no. The dimensions parameter is a tuple of elements in
    the order of (unit, top, bottom, left, right).  
    """
    if pg_no in list(range(len(document.sections))):
        pass
    else:
        print('The arguement passed into pg_no is not valid')
        return
    if isinstance(dimensions, tuple):
        page = document.sections[pg_no]
        page.top_margin = dimensions[0](dimensions[1])
        page.bottom_margin = dimensions[0](dimensions[2])
        page.left_margin = dimensions[0](dimensions[3])
        page.right_margin = dimensions[0](dimensions[4])  
    else:
        print('The dimensions parameter requires a tuple (unit, top, bottom, left, right)')

#change_margins(document, (Cm, 0.8, 0.8, 1.9, 1.9), 1)







        
def add_page(x):
    """
    This function adds x number of pages to the existing pages
    in the docx file.
    """
    for num in range(x):
        new_section = document.add_section(WD_SECTION.NEW_PAGE)

#add_page(2)










def word_table(document, dictionary_data, columns_per_table, table_style = 'Table Normal', row_info = [], row_headers = False, compact_text = True):

    data_categories = list(dictionary_data.keys())
    category_content = list(dictionary_data.values())

    additional_pages = int(ceil(len(dictionary_data)/columns_per_table - 1)) #number of category-distractions / columns_per_table -1  ==  10 / 5 - 1  ==  1
    total_docx_pages = additional_pages + 1

    if row_headers == True:
        columns_per_table = columns_per_table + 1  # uncomment if row labels required

    maxi = max([len(item) for item in category_content]) #largest number of rows required by longest list of triggers of a category
    header_row = 1
    total_rows = maxi + header_row

    if additional_pages > 0: #IF MORE THAN 1 PAGE REQUIRED FOR TABLE.
        for page in range(total_docx_pages):
            display_board = document.add_table(rows = total_rows, cols = columns_per_table, style = 'Table Grid') #rows = total_rows
            display_board = document.tables[page]
            if table_style != 'Table Normal':
                display_board.style = table_style

            if row_headers == True:
                all_columns = list(range(1, columns_per_table))
            else:
                all_columns = range(columns_per_table)
            
            headers = display_board.rows[0].cells
            for column in all_columns: 
                if data_categories != []:
                    data_header = data_categories.pop(0)
                    headers[column].text = data_header.upper()
                else:
                    break
              
            for col in all_columns: 
                each_data_sect = display_board.columns[col].cells
                if category_content != []:
                    content = category_content.pop(0)
                    for item, row_num in zip(content, list(range(maxi))):
                        if compact_text == True:
                            each_data_sect[row_num + 1].text = '\n' + item + '\n'
                        else:
                            each_data_sect[row_num + 1].text = item 
                else:
                    break
                
            if row_headers == True:
                column_0 = row_info #pass in a list
                row_labels = display_board.columns[0].cells
                for item, row in zip(column_0, range(total_rows)):
                    row_labels[row].text = item

                
            if page == additional_pages:
                break
            document.add_page_break() #new page added


    elif additional_pages == 0: #IF ONLY 1 PAGE REQUIRED FOR TABLE.
        display_board = document.add_table(rows= total_rows, cols = columns_per_table, style = 'Table Grid')
        if table_style != 'Table Normal':
            display_board.style = table_style

        if row_headers == True:
            all_columns = list(range(1, columns_per_table))
        else:
            all_columns = range(columns_per_table)

        headers = display_board.rows[0].cells
        for column in all_columns:  
            if data_categories != []:
                data_header = data_categories.pop(0)
                headers[column].text = data_header.upper()
            else:
                break

        for col in all_columns:
            each_data_sect = display_board.columns[col].cells
            if category_content != []:
                content = category_content.pop(0)                
                for item, row_num in zip(content, list(range(maxi))):
                    if compact_text == True:
                        each_data_sect[row_num + 1].text = '\n' + item + '\n'
                    else:
                        each_data_sect[row_num + 1].text = item 
            else:
                break
            
        if row_headers == True:
            column_0 = row_info #pass in a list
            row_labels = display_board.columns[0].cells
            for item, row in zip(column_0, range(total_rows)):
                row_labels[row].text = item
 

#The following are a few table styles:
#Colorful Grid
#Colorful Grid Accent 1
#Colorful Grid Accent 2
#Colorful Grid Accent 3
#Colorful Grid Accent 4
#Colorful Grid Accent 5
#Colorful Grid Accent 6
#Colorful List


#TABLE SAMPLE CODE-RUN:
lg_sample_dict = {'a': ['1', '2', '3'], 'b': ['1', '2', '3'], 'c': ['1', '2', '3'], 'd': ['1', '2', '3'],
                  'e': ['1', '2', '3'], 'f': ['1', '2', '3'], 'g': ['1', '2', '3'], 'h': ['1', '2', '3']}
sh_sample_dict = {'a': ['1', '2', '3'], 'b': ['1', '2', '3']}
headers = ['H1', 'h2', 'h3', 'h4']
#page_orientation_size(document, (Cm, 29.7, 42), 0, False)
#change_margins(document, (Cm, 1, 1, 1, 1), 0)

#word_table(sh_sample_dict, 2) 
#word_table(sh_sample_dict, 2, 'Light Grid', headers, True) 

#word_table(lg_sample_dict, 4)
#word_table(lg_sample_dict, 4, 'Light Grid', headers, True)









#document.save("WordAutoTestDocument.docx")




















