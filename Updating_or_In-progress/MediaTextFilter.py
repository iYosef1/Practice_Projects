import os
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from moviepy.editor import *

from ibm_watson import SpeechToTextV1  # module must be reactivated via new api key & url, if inactive for 30 consecutive days
from ibm_cloud_sdk_core.authenticators import IAMAuthenticator  # https://cloud.ibm.com/catalog#services 

import cv2
import pytesseract #https://github.com/UB-Mannheim/tesseract/wiki
pytesseract.pytesseract.tesseract_cmd = 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'



def txt_from_image(img_name, img_size = 3, show_image = False, error_notice = False, create_docx = True):
    """
    This function scrapes the text from an image and returns the text as strings.
    The parameters for this function include image's name & it's extension. The
    size can be readjusted for the accuracy of the scraped text but the preset size
    of 3 works fairly well. By default, all text scraped is compiled into a docx file.
    """
    split_tup = os.path.splitext(img_name)
    if split_tup[1] in ['.png', '.PNG', '.jpeg', '.JPEG', '.jpg', '.JPG']:
        img = img_name
    else:
        if error_notice == True:
            print('Enter appropriate file name and/or extension.')
        return
    img = cv2.imread(img, 0)
    img = cv2.resize(img, (0, 0), fx = img_size, fy = img_size) #proportional size change
    threshold_value = 169
    _, th1 = cv2.threshold(img, threshold_value, 255, cv2.THRESH_BINARY)
    text = pytesseract.image_to_string(th1)

    if create_docx == True:
        directory_content = os.listdir(os.getcwd())
        if 'ImageTranscript.docx' not in directory_content:
            word_file = docx.Document()
        else:
            word_file = docx.Document('ImageTranscript.docx')
            
        transcript_heading = 'IMAGE - FILENAME:  ' + split_tup[0]    
        transcript_heading = word_file.add_paragraph(transcript_heading)
        transcript_heading.runs[0].bold = True
        transcript_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        image_text = text.strip().split('\n')
        for text_line in image_text:
            paragraph = word_file.add_paragraph(text_line)
            paragraph_format = paragraph.paragraph_format  
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing = Pt(0)
        word_file.add_paragraph('======================================================================')
        word_file.save('ImageTranscript.docx')
    if show_image == True:
        show_img = cv2.imshow('original image', img)
        show_updated_img = cv2.imshow('updated', th1)

        cv2.waitKey(0)
        cv2.destroyAllWindows()
    return text

#print(txt_from_image('pg1.png'))
#print(txt_from_image('1.png'))







#ibm watson lite plan must be renewed if inactive for 30 days

def audio_transcript(media_file, create_docx = True):
    """
    This function transcribes the text from any video/audio extension file with audio.
    The transcript will be printed in IDLE. The create_docx parameter is set to True by
    default to create a word file in the directory of the this .py file. 
    """
    mp3_created = False
    split_tup = os.path.splitext(media_file)
    if split_tup[1] in ['.m4a', '.wav', '.avi', '.flv', '.mov', '.wmv', '.qt', '.mkv', '.amv', '.mp4', '.mpeg', '.m4v']:
        mp3_created = True
        mp3_file = 'converted_' + split_tup[0] + '.mp3'
        try:
            audio_file = AudioFileClip(media_file)
            audio_file.write_audiofile(mp3_file)
            audio_file.close()
        except:
            video_file = VideoFileClip(media_file)
            extracted_audio = video_file.audio
            extracted_audio.write_audiofile(mp3_file)
            extracted_audio.close()
            video_file.close()

    elif split_tup[1] == '.mp3':
        mp3_file = media_file

    else:
        return

    api_key = '5dmldad256RsjIzahMo56PPpDRsImOCydolHJqX4dKo3'
    url = 'https://api.us-south.speech-to-text.watson.cloud.ibm.com/instances/fd5b56ae-e50a-4e85-a5f6-d335d3a8b312'

    authenticator = IAMAuthenticator(api_key)
    stt = SpeechToTextV1(authenticator = authenticator)
    stt. set_service_url(url)

    with open(mp3_file, 'rb') as f:
        res = stt.recognize(audio = f, content_type = 'audio/mp3', model = 'en-US_NarrowbandModel', continuous = True).get_result()

    if mp3_created == True:
        os.unlink(mp3_file)

    for statement in res['results']:
        print(statement['alternatives'][0])
    print('----------------------------')
        
    if create_docx == True:
        directory_content = os.listdir(os.getcwd())
        if 'AudioTranscript.docx' not in directory_content:
            word_file = docx.Document()
        else:
            word_file = docx.Document('AudioTranscript.docx')
        heading_per_transcript = 'AUDIO - FILENAME:  ' + split_tup[0]
        heading_per_transcript = word_file.add_paragraph(heading_per_transcript)
        heading_per_transcript.runs[0].bold = True
        heading_per_transcript.alignment = WD_ALIGN_PARAGRAPH.CENTER
        counter = 1
        for statement in res['results']:
            label = f'Statement {str(counter)}: '
            label = word_file.add_paragraph(label)
            label.runs[0].bold = True
            statement = statement['alternatives'][0]['transcript'] + ',  ' + 'confidence level:  ' + str(statement['alternatives'][0]['confidence'])
            label.add_run(statement)
            counter += 1
        word_file.add_paragraph('======================================================================')
        word_file.save('AudioTranscript.docx')


#audio_transcript('test1.mp3')
#audio_transcript('test2.mp3')

#results = func_on_dir(audio_transcript, target = True)














def func_on_dir(funcs_tup, directory = None, target = False, dir_check = False):
    """
    This function runs through every folder and/or file in it's cwd and applies the
    function(s) in the tuple parameter, funcs_tup. The 1st function will be applied
    on files, the 2nd function on folders. If only 1 function is passed into parameter,
    then only 1 function will run on either files alone (target = True) or both files
    and folders (target = False). If a directory consists of both files & folders BUT
    a function is required to be performed on ONLY folders, then pass in a null function
    for files (target = True).
    """
    if directory == None:
        cwd = os.getcwd()
    else:
        cwd = os.chdir(directory)

    directory_content = os.listdir(cwd)
    #print(directory_content)
    results = []
    if target == True:
        for item in directory_content:
            file_check = os.path.isfile(item)
            if file_check == True:
                if type(funcs_tup) != type(()):
                    action = (item, funcs_tup(item))   
                elif type(funcs_tup) == type(()):
                    action = (item, funcs_tup[0](item))
            elif file_check == False:
                if type(funcs_tup) == type(()):
                    action = (item, funcs_tup[1](item))
                else:
                    action = (item, None)
            results.append(action)
    else:
        for item in directory_content:
            action = (item, funcs_tup(item))
            results.append(action)
    return results

#results = func_on_dir(funcs_tup = (txt_from_image), target = True)
#print(results)





'''
def null_func(x):
    pass

pt_switch = True
def print_switch(item, new):
    if pt_switch == True:
        print(item, new)
    else:
        pass


def dir_walk_transform(funcs = (null_func, null_func), pt_switch = True):
    """
    This function transforms every file and/or folder with the tupe of
    2 functions passed in as the parameter.
    """
    for folder_name, subfolders, filenames in os.walk('.\\'):
        if os.path.basename(folder_name) in os.listdir('.\\'):
            folder_name = folder_name.upper()
            print_switch('\n\n')
            print_switch('XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX')
            print_switch('XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX')
        print_switch('Folder: ', folder_name)
        
        if subfolders != []:
            print_switch('Subfolders in', folder_name, 'are:', subfolders)
            for folder in subfolders:
                print_switch('\n------------------------------------------------------------------------------------------------')
                print_switch('FOLDER NAME:', folder)
                print_switch(funcs[0](folder_name + '\\' + folder))
                print_switch('------------------------------------------------------------------------------------------------\n')
        else:
            print_switch('Subfolders: N/A')
            
        if filenames != []:
            print_switch('Filenames in', folder_name, 'are:', filenames)
            for file in filenames:
                print_switch('\n------------------------------------------------------------------------------------------------')
                print_switch('FILE NAME:', file)
                print_switch(funcs[1](folder_name + '\\' + file))
                print_switch('------------------------------------------------------------------------------------------------\n')
        else:
            print_switch('Files: N/A')
            
        print_switch('')
        
    print_switch('\n\n\nFINISHED.')

dir_walk_transform(funcs = (null_func, txt_from_image))
'''

    




def null_func(x):
    pass



def dir_walk_transform(funcs = (null_func, null_func)):
    """
    This function transforms every file and/or folder with the tupe of
    2 functions passed in as the parameter.
    """
    for folder_name, subfolders, filenames in os.walk('.\\'):
        if os.path.basename(folder_name) in os.listdir('.\\'):
            folder_name = folder_name.upper()
            print('\n\n')
            print('XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX')
            print('XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX')
        print('Folder: ', folder_name)
        
        if subfolders != []:
            print('Subfolders in', folder_name, 'are:', subfolders)
            for folder in subfolders:
                print('\n------------------------------------------------------------------------------------------------')
                print('FOLDER NAME:', folder)
                print(funcs[0](folder_name + '\\' + folder))
                print('------------------------------------------------------------------------------------------------\n')
        else:
            print('Subfolders: N/A')
            
        if filenames != []:
            print('Filenames in', folder_name, 'are:', filenames)
            for file in filenames:
                print('\n------------------------------------------------------------------------------------------------')
                print('FILE NAME:', file)
                print(funcs[1](folder_name + '\\' + file))
                print('------------------------------------------------------------------------------------------------\n')
        else:
            print('Files: N/A')
            
        print('')
        
    print('\n\n\nFINISHED.')

#dir_walk_transform(funcs = (null_func, txt_from_image))



